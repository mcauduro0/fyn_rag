"""
DOCX Document Processor.
Extracts text content from Microsoft Word documents.
"""

import logging
from typing import Dict, Any
from docx import Document

from app.data.processors.base_processor import BaseProcessor

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):
    """Processor for DOCX documents."""
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Validate file
            path = self._validate_file(file_path, ['.docx', '.DOCX'])
            
            logger.info(f"Processing DOCX: {path}")
            
            # Read DOCX
            doc = Document(str(path))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(" | ".join(row_data))
                
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            
            if tables_text:
                full_text += "\n\n--- Tables ---\n\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
                "file_name": path.name,
                "file_size_bytes": path.stat().st_size,
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            logger.info(f"Successfully extracted {len(full_text)} characters from {metadata['num_paragraphs']} paragraphs")
            
            return self._success_response(full_text, metadata, path)
            
        except Exception as e:
            return self._error_response(e, file_path)


if __name__ == "__main__":
    # Test the DOCX processor
    logging.basicConfig(level=logging.INFO)
    
    processor = DOCXProcessor()
    
    # Test with a sample DOCX (you'll need to provide a path)
    test_file = "sample.docx"
    
    result = processor.process(test_file)
    
    if result["success"]:
        print(f"\nSuccessfully processed DOCX:")
        print(f"Paragraphs: {result['metadata']['num_paragraphs']}")
        print(f"Tables: {result['metadata']['num_tables']}")
        print(f"Content length: {result['content_length']} characters")
        print(f"\nFirst 500 characters:")
        print(result["content"][:500])
    else:
        print(f"\nError: {result['error']}")
